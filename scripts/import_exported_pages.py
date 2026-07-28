#!/usr/bin/env python3
"""
Aastha bulk importer — DOCX-only fallback edition.

What it accepts:
1. DOCX + JSON-LD pair: uses the supplied JSON-LD.
2. DOCX without JSON-LD: imports the visible content and generates a safe,
   minimal JSON-LD graph automatically.
3. Files may be in any order or subfolder under content-drop/.

Important:
- Supplied full JSON-LD remains preferable because it may contain richer
  FAQPage, MedicalCondition and MedicalProcedure entities.
- Generated fallback schema includes the page, doctor, both clinics and
  breadcrumbs. It is valid but intentionally conservative.
"""

from __future__ import annotations

import argparse
import csv
import html
import json
import re
import shutil
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path
from typing import Optional
from urllib.parse import unquote, urlparse

try:
    from docx import Document
except ImportError as exc:
    raise SystemExit(
        "python-docx is unavailable in this Python environment. "
        "Run 0_Test_Python_Environment.bat first."
    ) from exc

SITE_ROOT = Path(__file__).resolve().parents[1]
DROP_DIR = SITE_ROOT / "content-drop"
REPORT_DIR = SITE_ROOT / "reports"
BACKUP_DIR = SITE_ROOT / "backups"
SITEMAP_PATH = SITE_ROOT / "sitemap.xml"
PRIMARY_ORIGIN = "https://www.aasthaskincentre.in"

DISCLAIMER = (
    "Results may vary from person to person. This content is for informational purposes only "
    "and does not substitute a consultation with a qualified medical professional. "
    "Dr. Cheena Langer, MD is a registered medical practitioner."
)

# URLs used in older DOCX exports that differ from the current locked sitemap.
LEGACY_PATH_MAP = {
    "/patient-concerns/": "/conditions/",
    "/treatments/botox-dermal-fillers/": "/treatments/botulinum-toxin-dermal-fillers/",
    "/treatments/eczema-treatment/": "/treatments/eczema-atopic-dermatitis-treatment/",
    "/treatments/black-neck-acanthosis-nigricans/": "/treatments/black-neck-acanthosis-nigricans-treatment/",
}

# Filename fallbacks for core pages whose DOCX metadata is not detected.
CORE_FILENAME_MAP = {
    "homepage": "/",
    "aboutdrcheena": "/dr-cheena-langer/",
    "drcheena": "/dr-cheena-langer/",
    "aboutaastha": "/about/",
    "patientconcernsoverviewpage": "/conditions/",
    "patientconcerns": "/conditions/",
    "treatmentsoverviewpage": "/treatments/",
    "treatmentsoverview": "/treatments/",
    "karannagarjammu": "/locations/karan-nagar/",
    "karannagar": "/locations/karan-nagar/",
    "paloura": "/locations/paloura/",
    "appointment": "/book-appointment/",
    "contact": "/contact/",
}

SKIP_SECTION_NAMES = {
    "meta title",
    "meta description",
    "recommended url",
    "canonical url",
    "h1",
    "geo entity summary",
    "json-ld schema",
}

def normalise_label(text: str) -> str:
    text = re.sub(r"^#{1,6}\s*", "", text.strip())
    text = text.strip("*_` ")
    text = re.sub(r"^\d+\s*[.)-]\s*", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip().lower()

def normalise_filename(text: str) -> str:
    text = unquote(text)
    text = re.sub(r"\s*-\s*json[- ]?ld$", "", text, flags=re.I)
    text = re.sub(r"^\d+\s*[.)_-]*\s*", "", text)
    return re.sub(r"[^a-z0-9]+", "", text.lower())

def clean_text(text: str) -> str:
    text = re.sub(r"(?:cite|filecite).*?", "", text)
    return re.sub(r"\s+", " ", text).strip()

def normalise_path(value: str) -> str:
    value = unquote(value.strip())
    parsed = urlparse(value)
    path = parsed.path if parsed.scheme or parsed.netloc else value
    if not path.startswith("/"):
        path = "/" + path
    if not path.endswith("/"):
        path += "/"
    path = re.sub(r"/{2,}", "/", path)
    return LEGACY_PATH_MAP.get(path, path)

def canonical_url(path: str) -> str:
    return PRIMARY_ORIGIN + normalise_path(path)

def iter_docx_paragraphs_and_tables(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph.text, paragraph.style.name if paragraph.style else ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph.text, paragraph.style.name if paragraph.style else ""

def paragraph_entries(docx_path: Path) -> list[dict]:
    doc = Document(docx_path)
    entries = []
    for raw_text, style in iter_docx_paragraphs_and_tables(doc):
        text = clean_text(raw_text)
        if text:
            entries.append({"text": text, "style": style})
    return entries

def next_value(entries: list[dict], section_name: str) -> Optional[str]:
    target = normalise_label(section_name)
    for index, entry in enumerate(entries):
        if normalise_label(entry["text"]) == target:
            for following in entries[index + 1:]:
                value = following["text"].strip("*_` ")
                if value:
                    return value
    return None

def extract_url_anywhere(entries: list[dict]) -> Optional[str]:
    labelled = next_value(entries, "Recommended URL") or next_value(entries, "Canonical URL")
    if labelled:
        return normalise_path(labelled)

    for entry in entries:
        text = unquote(entry["text"])
        match = re.search(r"https?://(?:www\.)?aasthaskincentre\.(?:in|com)(/[^\s]*)?", text, flags=re.I)
        if match:
            return normalise_path(match.group(1) or "/")

    for entry in entries:
        text = entry["text"].strip()
        if re.fullmatch(r"/[a-z0-9][a-z0-9/_-]*/?", text, flags=re.I):
            return normalise_path(text)
    return None

def infer_core_path_from_filename(path: Path) -> Optional[str]:
    key = normalise_filename(path.stem)
    return CORE_FILENAME_MAP.get(key)

def json_page_path(data: object) -> str:
    if not isinstance(data, dict):
        raise ValueError("JSON-LD root must be an object.")
    graph = data.get("@graph", [data])
    if not isinstance(graph, list):
        graph = [graph]
    urls = []
    for item in graph:
        if not isinstance(item, dict):
            continue
        if item.get("@type") in {"MedicalWebPage", "WebPage", "CollectionPage", "AboutPage", "ProfilePage", "ContactPage"} and item.get("url"):
            urls.append(str(item["url"]))
    if not urls:
        for item in graph:
            if isinstance(item, dict) and item.get("url"):
                urls.append(str(item["url"]))
    if not urls:
        raise ValueError("No page URL found in JSON-LD.")
    return normalise_path(urls[0])

def escape_inline(text: str) -> str:
    value = html.escape(text, quote=False)
    value = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", value)
    value = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<em>\1</em>", value)
    value = re.sub(r"`([^`]+)`", r"<code>\1</code>", value)
    return value

def entries_to_html(entries: list[dict]) -> tuple[str, str]:
    h1 = next_value(entries, "H1") or "Dermatology Care in Jammu"

    start_index = None
    for index, entry in enumerate(entries):
        if normalise_label(entry["text"]) == "hero section":
            start_index = index + 1
            break
    if start_index is None:
        for index, entry in enumerate(entries):
            if normalise_label(entry["text"]) == "geo entity summary":
                start_index = min(index + 2, len(entries))
                break
    if start_index is None:
        # Start after H1 value when possible.
        h1_value = next_value(entries, "H1")
        if h1_value:
            for index, entry in enumerate(entries):
                if entry["text"] == h1_value:
                    start_index = index + 1
                    break
    if start_index is None:
        start_index = 0

    output = []
    list_mode = None

    def close_list():
        nonlocal list_mode
        if list_mode:
            output.append(f"</{list_mode}>")
            list_mode = None

    for entry in entries[start_index:]:
        text = entry["text"]
        norm = normalise_label(text)
        style = entry["style"].lower()

        if norm in SKIP_SECTION_NAMES:
            continue

        if text == DISCLAIMER:
            close_list()
            output.append(
                '<section class="medical-disclaimer notice">'
                '<strong>Medical disclaimer:</strong> ' + html.escape(text) + "</section>"
            )
            break

        if norm in {
            "primary buttons",
            "supporting line",
            "suggested internal link",
            "suggested internal links",
            "appointment wording",
            "important availability wording",
        }:
            close_list()
            continue

        if "heading" in style:
            close_list()
            cleaned = re.sub(r"^\d+\s*[.)-]\s*", "", text).strip()
            if normalise_label(cleaned) in SKIP_SECTION_NAMES:
                continue
            level = 3 if ("heading 3" in style or "heading 4" in style) else 2
            output.append(f"<h{level}>{escape_inline(cleaned)}</h{level}>")
            continue

        if style.startswith("list bullet"):
            if list_mode != "ul":
                close_list()
                output.append("<ul>")
                list_mode = "ul"
            output.append(f"<li>{escape_inline(text)}</li>")
            continue

        if style.startswith("list number"):
            if list_mode != "ol":
                close_list()
                output.append("<ol>")
                list_mode = "ol"
            output.append(f"<li>{escape_inline(text)}</li>")
            continue

        close_list()
        output.append(f"<p>{escape_inline(text)}</p>")

    close_list()
    return h1.strip("*_` "), "\n".join(output)

def replace_one(source: str, pattern: str, replacement: str) -> str:
    updated, count = re.subn(pattern, replacement, source, count=1, flags=re.I | re.S)
    if count == 0:
        raise ValueError(f"Expected HTML pattern not found: {pattern}")
    return updated

def page_type_for_path(path: str) -> str:
    if path == "/":
        return "MedicalWebPage"
    if path == "/about/":
        return "AboutPage"
    if path == "/dr-cheena-langer/":
        return "ProfilePage"
    if path in {"/conditions/", "/treatments/"}:
        return "CollectionPage"
    if path == "/contact/":
        return "ContactPage"
    return "MedicalWebPage"

def breadcrumb_items(path: str, h1: str) -> list[dict]:
    items = [
        {"@type": "ListItem", "position": 1, "name": "Home", "item": PRIMARY_ORIGIN + "/"}
    ]
    if path.startswith("/treatments/") and path != "/treatments/":
        items.append({
            "@type": "ListItem",
            "position": 2,
            "name": "Treatments",
            "item": PRIMARY_ORIGIN + "/treatments/",
        })
    elif path.startswith("/locations/"):
        items.append({
            "@type": "ListItem",
            "position": 2,
            "name": "Locations",
            "item": PRIMARY_ORIGIN + "/contact/",
        })
    if path != "/":
        items.append({
            "@type": "ListItem",
            "position": len(items) + 1,
            "name": h1,
            "item": canonical_url(path),
        })
    return items

def generate_minimal_schema(path: str, title: str, description: str, h1: str) -> dict:
    page_url = canonical_url(path)
    page_type = page_type_for_path(path)

    graph = [
        {
            "@type": page_type,
            "@id": page_url + "#webpage",
            "url": page_url,
            "name": title,
            "description": description,
            "isPartOf": {"@id": PRIMARY_ORIGIN + "/#website"},
            "reviewedBy": {"@id": PRIMARY_ORIGIN + "/dr-cheena-langer/#physician"},
            "breadcrumb": {"@id": page_url + "#breadcrumb"},
            "inLanguage": "en-IN",
        },
        {
            "@type": "Physician",
            "@id": PRIMARY_ORIGIN + "/dr-cheena-langer/#physician",
            "name": "Dr. Cheena Langer",
            "honorificSuffix": "MBBS, MD Dermatology",
            "jobTitle": "Consultant Dermatologist",
            "url": PRIMARY_ORIGIN + "/dr-cheena-langer/",
            "telephone": ["+91-7006613362", "+91-9796676541"],
            "workLocation": [
                {"@id": PRIMARY_ORIGIN + "/locations/karan-nagar/#clinic"},
                {"@id": PRIMARY_ORIGIN + "/locations/paloura/#clinic"},
            ],
        },
        {
            "@type": "MedicalClinic",
            "@id": PRIMARY_ORIGIN + "/locations/karan-nagar/#clinic",
            "name": "Aastha Skin & Dermato-Cosmetic Centre - Karan Nagar",
            "url": PRIMARY_ORIGIN + "/locations/karan-nagar/",
            "telephone": ["+91-7006613362", "+91-9796676541", "+91-191-3509230"],
            "address": {
                "@type": "PostalAddress",
                "streetAddress": "Lane 2, Karan Nagar, near Amphalla Chowk",
                "addressLocality": "Jammu",
                "addressRegion": "Jammu and Kashmir",
                "postalCode": "180005",
                "addressCountry": "IN",
            },
        },
        {
            "@type": "MedicalClinic",
            "@id": PRIMARY_ORIGIN + "/locations/paloura/#clinic",
            "name": "Aastha Skin & Dermato-Cosmetic Centre - Paloura",
            "url": PRIMARY_ORIGIN + "/locations/paloura/",
            "telephone": ["+91-7006613362", "+91-9796676541", "+91-191-3135864"],
            "address": {
                "@type": "PostalAddress",
                "streetAddress": "Paloura Chowk, Top Paloura, opposite Government Senior Secondary School",
                "addressLocality": "Jammu",
                "addressRegion": "Jammu and Kashmir",
                "postalCode": "181121",
                "addressCountry": "IN",
            },
        },
        {
            "@type": "BreadcrumbList",
            "@id": page_url + "#breadcrumb",
            "itemListElement": breadcrumb_items(path, h1),
        },
    ]

    if path == "/":
        graph.insert(0, {
            "@type": "WebSite",
            "@id": PRIMARY_ORIGIN + "/#website",
            "url": PRIMARY_ORIGIN + "/",
            "name": "Aastha Skin & Dermato-Cosmetic Centre",
            "inLanguage": "en-IN",
        })

    return {"@context": "https://schema.org", "@graph": graph}

def build_main(h1: str, summary: str, body_html: str, path: str) -> str:
    parent_label = "Treatments" if path.startswith("/treatments/") else "Home"
    parent_url = "/treatments/" if path.startswith("/treatments/") else "/"
    breadcrumb_mid = (
        f' / <a href="{parent_url}">{parent_label}</a>'
        if path != "/" else ""
    )
    return f"""<main id="main-content">
<section class="page-hero"><div class="container">
  <nav class="breadcrumbs" aria-label="Breadcrumb"><a href="/">Home</a>{breadcrumb_mid} / <span aria-current="page">{html.escape(h1)}</span></nav>
  <span class="eyebrow">Dermatologist-led care in Jammu</span>
  <h1>{html.escape(h1)}</h1>
  <p class="lead">{html.escape(summary)}</p>
  <div class="hero-actions">
    <a class="button" href="/book-appointment/">Book an appointment</a>
    <a class="button button-secondary" href="https://wa.me/917006613362" target="_blank" rel="noopener">WhatsApp the clinic</a>
  </div>
</div></section>
<section class="section"><div class="container content-layout">
  <article class="prose">
    {body_html}
  </article>
  <aside class="sidebar-card">
    <h2>Book a consultation</h2>
    <p>Consult Dr. Cheena Langer at Karan Nagar or Paloura Chowk.</p>
    <p><strong>Consultation fee:</strong> ₹500</p>
    <a class="button" href="/book-appointment/">Request appointment</a>
    <a class="button button-secondary" href="tel:+917006613362">Call 7006613362</a>
  </aside>
</div></section>
</main>"""

def destination_for_path(path: str) -> Path:
    if path == "/":
        return SITE_ROOT / "index.html"
    return SITE_ROOT / path.strip("/") / "index.html"

def scan_docx() -> list[dict]:
    records = []
    for path in sorted(DROP_DIR.rglob("*.docx")):
        try:
            entries = paragraph_entries(path)
            page_path = extract_url_anywhere(entries) or infer_core_path_from_filename(path)
            records.append({
                "path": path,
                "entries": entries,
                "page_path": page_path,
                "name_key": normalise_filename(path.stem),
                "error": "",
            })
        except Exception as exc:
            records.append({
                "path": path,
                "entries": [],
                "page_path": None,
                "name_key": normalise_filename(path.stem),
                "error": str(exc),
            })
    return records

def scan_json() -> list[dict]:
    records = []
    for path in sorted(DROP_DIR.rglob("*.json")):
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            records.append({
                "path": path,
                "data": data,
                "page_path": json_page_path(data),
                "name_key": normalise_filename(path.stem),
                "error": "",
            })
        except Exception as exc:
            records.append({
                "path": path,
                "data": None,
                "page_path": None,
                "name_key": normalise_filename(path.stem),
                "error": str(exc),
            })
    return records

def pair_records(docx_records: list[dict], json_records: list[dict]):
    json_by_url = defaultdict(list)
    json_by_name = defaultdict(list)
    issues = []

    for record in json_records:
        if record["error"]:
            issues.append({
                "status": "ERROR",
                "page": "",
                "docx": "",
                "json": str(record["path"]),
                "message": record["error"],
            })
            continue
        json_by_url[record["page_path"]].append(record)
        json_by_name[record["name_key"]].append(record)

    used_json = set()
    pairs = []

    for doc in docx_records:
        if doc["error"]:
            issues.append({
                "status": "ERROR",
                "page": "",
                "docx": str(doc["path"]),
                "json": "",
                "message": doc["error"],
            })
            continue

        if not doc["page_path"]:
            issues.append({
                "status": "ERROR",
                "page": "",
                "docx": str(doc["path"]),
                "json": "",
                "message": "Could not determine the destination URL from DOCX content or filename.",
            })
            continue

        candidates = [j for j in json_by_url.get(doc["page_path"], []) if j["path"] not in used_json]
        method = "URL"

        if not candidates:
            candidates = [j for j in json_by_name.get(doc["name_key"], []) if j["path"] not in used_json]
            method = "filename"

        if len(candidates) > 1:
            issues.append({
                "status": "ERROR",
                "page": doc["page_path"],
                "docx": str(doc["path"]),
                "json": " | ".join(str(j["path"]) for j in candidates),
                "message": "Multiple JSON-LD files match this DOCX. Remove duplicates.",
            })
            continue

        js = candidates[0] if candidates else None
        if js:
            used_json.add(js["path"])
            method = "supplied JSON by " + method
        else:
            method = "generated minimal JSON-LD"

        pairs.append({"doc": doc, "json": js, "method": method})

    for js in json_records:
        if js["error"] or js["path"] in used_json:
            continue
        issues.append({
            "status": "WARNING",
            "page": js["page_path"] or "",
            "docx": "",
            "json": str(js["path"]),
            "message": "JSON-LD was not paired with a DOCX file.",
        })

    return pairs, issues

def validate_pair(pair: dict) -> list[str]:
    errors = []
    entries = pair["doc"]["entries"]
    path = pair["doc"]["page_path"]
    title = next_value(entries, "Meta title")
    description = next_value(entries, "Meta description")
    h1 = next_value(entries, "H1")

    if not title:
        errors.append("Missing Meta title section/value.")
    elif len(title) > 63:
        errors.append(f"Meta title is {len(title)} characters; shorten it before launch.")

    if not description:
        errors.append("Missing Meta description section/value.")
    elif len(description) > 160:
        errors.append(f"Meta description is {len(description)} characters; shorten it before launch.")

    if not h1:
        errors.append("Missing H1 section/value.")

    destination = destination_for_path(path)
    if not destination.exists():
        errors.append(f"No matching HTML page exists for {path}.")

    return errors

def create_backup(destination: Path, backup_root: Path):
    target = backup_root / destination.relative_to(SITE_ROOT)
    target.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(destination, target)

def import_pair(pair: dict, backup_root: Path, dry_run: bool) -> dict:
    errors = validate_pair(pair)
    doc = pair["doc"]
    path = doc["page_path"]

    if errors:
        return {
            "status": "ERROR",
            "page": path,
            "docx": str(doc["path"]),
            "json": str(pair["json"]["path"]) if pair["json"] else "",
            "message": " ".join(errors),
        }

    entries = doc["entries"]
    title = next_value(entries, "Meta title")
    description = next_value(entries, "Meta description")
    summary = next_value(entries, "GEO entity summary") or (
        "Dr. Cheena Langer is an MD dermatologist leading Aastha Skin & Dermato-Cosmetic Centre in Jammu. "
        "The clinic provides advanced medical and aesthetic skin treatments across two locations at Karan Nagar and Paloura Chowk."
    )
    h1, body_html = entries_to_html(entries)
    destination = destination_for_path(path)

    if pair["json"]:
        schema_data = pair["json"]["data"]
        schema_source = "supplied full JSON-LD"
    else:
        schema_data = generate_minimal_schema(path, title, description, h1)
        schema_source = "generated minimal JSON-LD"

    if not dry_run:
        create_backup(destination, backup_root)
        current = destination.read_text(encoding="utf-8")
        current = replace_one(current, r"<title>.*?</title>", f"<title>{html.escape(title)}</title>")
        current = replace_one(
            current,
            r'<meta\s+name="description"\s+content=".*?">',
            f'<meta name="description" content="{html.escape(description, quote=True)}">'
        )
        current = replace_one(
            current,
            r'<meta\s+name="robots"\s+content=".*?">',
            '<meta name="robots" content="index,follow">'
        )
        current = replace_one(
            current,
            r'<link\s+rel="canonical"\s+href=".*?">',
            f'<link rel="canonical" href="{canonical_url(path)}">'
        )

        json_text = json.dumps(schema_data, ensure_ascii=False, indent=2)
        json_text = re.sub(
            r"https://(?:www\.)?aasthaskincentre\.(?:in|com)",
            PRIMARY_ORIGIN,
            json_text,
            flags=re.I,
        )
        schema = f'<script type="application/ld+json">\n{json_text}\n</script>'

        current = re.sub(
            r'\s*<script type="application/ld\+json">.*?</script>\s*</head>',
            "\n</head>",
            current,
            flags=re.I | re.S,
        )
        current = current.replace("</head>", f"  {schema}\n</head>", 1)
        current = replace_one(
            current,
            r'<main\s+id="main-content">.*?</main>',
            build_main(h1, summary, body_html, path),
        )
        destination.write_text(current, encoding="utf-8")

    return {
        "status": "DRY-RUN OK" if dry_run else "IMPORTED",
        "page": path,
        "docx": str(doc["path"]),
        "json": str(pair["json"]["path"]) if pair["json"] else "",
        "message": f"{pair['method']}; {schema_source}.",
    }

def regenerate_sitemap():
    urls = []
    for html_file in SITE_ROOT.rglob("*.html"):
        if any(part in {"backups", "reports", "content-drop"} for part in html_file.parts):
            continue
        text = html_file.read_text(encoding="utf-8", errors="ignore")
        robots = re.search(r'<meta\s+name="robots"\s+content="([^"]+)"', text, flags=re.I)
        if robots and "noindex" in robots.group(1).lower():
            continue
        canonical = re.search(r'<link\s+rel="canonical"\s+href="([^"]+)"', text, flags=re.I)
        if canonical and canonical.group(1).startswith(PRIMARY_ORIGIN):
            urls.append(canonical.group(1))
    lines = [
        '<?xml version="1.0" encoding="UTF-8"?>',
        '<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">',
    ]
    lines.extend(f"  <url><loc>{html.escape(url)}</loc></url>" for url in sorted(set(urls)))
    lines.append("</urlset>")
    SITEMAP_PATH.write_text("\n".join(lines), encoding="utf-8")

def write_reports(rows: list[dict], mode: str):
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    csv_path = REPORT_DIR / f"import-report-{mode}-{stamp}.csv"
    html_path = REPORT_DIR / f"import-report-{mode}-{stamp}.html"
    fields = ["status", "page", "docx", "json", "message"]

    with csv_path.open("w", newline="", encoding="utf-8-sig") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)

    counts = defaultdict(int)
    for row in rows:
        counts[row["status"]] += 1

    table_rows = "\n".join(
        "<tr>" + "".join(
            f"<td>{html.escape(str(row.get(field, '')))}</td>" for field in fields
        ) + "</tr>"
        for row in rows
    )
    html_path.write_text(
        f"""<!doctype html>
<html lang="en"><head><meta charset="utf-8"><title>Aastha Import Report</title>
<style>
body{{font-family:Arial,sans-serif;margin:32px;color:#242124}}
h1{{color:#7b1e3a}}
.summary{{display:flex;gap:12px;flex-wrap:wrap;margin:20px 0}}
.badge{{padding:8px 12px;border-radius:999px;background:#fff4cc;font-weight:bold}}
table{{width:100%;border-collapse:collapse;font-size:14px}}
th,td{{border:1px solid #ddd;padding:9px;vertical-align:top;text-align:left}}
th{{background:#7b1e3a;color:white}}
tr:nth-child(even){{background:#faf7f8}}
</style></head><body>
<h1>Aastha Bulk Import Report</h1>
<p>Mode: <strong>{html.escape(mode)}</strong></p>
<div class="summary">{''.join(f'<span class="badge">{html.escape(k)}: {v}</span>' for k,v in sorted(counts.items()))}</div>
<table><thead><tr>{''.join(f'<th>{field}</th>' for field in fields)}</tr></thead><tbody>{table_rows}</tbody></table>
</body></html>""",
        encoding="utf-8",
    )
    return csv_path, html_path

def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--dry-run", action="store_true")
    args = parser.parse_args()

    docs = scan_docx()
    schemas = scan_json()

    if not docs:
        print(f"No DOCX files found under {DROP_DIR}")
        return 1

    pairs, issues = pair_records(docs, schemas)
    stamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup_root = BACKUP_DIR / stamp
    rows = list(issues)

    for pair in pairs:
        rows.append(import_pair(pair, backup_root, args.dry_run))

    if not args.dry_run:
        regenerate_sitemap()

    mode = "dry-run" if args.dry_run else "import"
    csv_path, html_path = write_reports(rows, mode)

    for row in rows:
        print(f"{row['status']}: {row['page']} — {row['message']}")
    print(f"\nCSV report: {csv_path}")
    print(f"HTML report: {html_path}")

    errors = sum(1 for row in rows if row["status"] == "ERROR")
    return 2 if errors else 0

if __name__ == "__main__":
    raise SystemExit(main())
